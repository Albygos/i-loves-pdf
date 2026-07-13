import os
import uuid
import time
import zipfile
import shutil
import io
import threading
from flask import Flask, request, jsonify, send_file, send_from_directory, redirect, abort
from pypdf import PdfReader, PdfWriter
import fitz  # PyMuPDF
from PIL import Image
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

from pptx import Presentation
from pptx.util import Inches, Pt
import openpyxl
from reportlab.platypus import Table as RLTable, TableStyle as RLTableStyle, PageBreak
from reportlab.lib import colors
from deep_translator import GoogleTranslator

from pdf_tools import *



app = Flask(__name__)

# Setup directories inside workspace
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
PROCESSED_FOLDER = os.path.join(BASE_DIR, 'processed')

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['PROCESSED_FOLDER'] = PROCESSED_FOLDER
# Limit maximum upload to 50MB
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

def clean_old_files():
    """Deletes files in uploads and processed folders older than 15 minutes."""
    while True:
        now = time.time()
        for folder in [UPLOAD_FOLDER, PROCESSED_FOLDER]:
            if os.path.exists(folder):
                for filename in os.listdir(folder):
                    file_path = os.path.join(folder, filename)
                    try:
                        # Skip special system files
                        if filename.startswith('.'):
                            continue
                        if os.path.isfile(file_path) or os.path.islink(file_path):
                            if os.path.getmtime(file_path) < now - 900:  # 15 minutes
                                os.unlink(file_path)
                        elif os.path.isdir(file_path):
                            if os.path.getmtime(file_path) < now - 900:
                                shutil.rmtree(file_path)
                    except Exception as e:
                        print(f"Error cleaning file {file_path}: {e}")
        time.sleep(300)  # Check every 5 minutes

# Start the clean up thread as a daemon
cleanup_thread = threading.Thread(target=clean_old_files, daemon=True)
cleanup_thread.start()

from werkzeug.utils import secure_filename
import traceback

def get_unique_filename(ext):
    return f"{uuid.uuid4().hex}.{ext}"

@app.errorhandler(Exception)
def handle_global_exception(e):
    traceback.print_exc()
    if isinstance(e, ValueError):
        return jsonify({'error': str(e)}), 400
    return jsonify({'error': 'An internal server error occurred.'}), 500

def process_upload(request_files, expected_extensions, max_files=1):
    """
    Generic utility to validate and save uploaded files.
    expected_extensions should be a list like ['.pdf'] or ['.jpg', '.png']
    """
    if not request_files:
        raise ValueError('No files uploaded.')
        
    saved_paths = []
    
    # In some forms it's request.files.getlist('files'), in others request.files.get('file')
    # This standardizes to a list.
    if isinstance(request_files, list):
        files = request_files
    else:
        files = [request_files]
        
    if len(files) > max_files:
        raise ValueError(f'Too many files uploaded. Maximum is {max_files}.')
        
    for f in files:
        if f.filename == '':
            continue
            
        ext = os.path.splitext(f.filename)[1].lower()
        if ext not in expected_extensions:
            raise ValueError(f'Invalid file type. Expected one of {", ".join(expected_extensions)}')
            
        secure_name = secure_filename(f.filename)
        # We append a UUID to prevent collisions but keep the safe original name (if any)
        final_name = f"{uuid.uuid4().hex}_{secure_name}" if secure_name else get_unique_filename(ext[1:])
        
        path = os.path.join(UPLOAD_FOLDER, final_name)
        f.save(path)
        saved_paths.append(path)
        
    if not saved_paths:
        raise ValueError('No valid files were processed.')
        
    return saved_paths


SUPPORTED_LANGS = {
    'en': 'English',
    'es': 'Español',
    'fr': 'Français',
    'de': 'Deutsch',
    'pt': 'Português',
    'hi': 'हिन्दी'
}

SLUG_TO_FILE = {
    '': 'index.html',
    'merge_pdf': 'merge.html',
    'split_pdf': 'split.html',
    'compress_pdf': 'compress.html',
    'word_to_pdf': 'word-to-pdf.html',
    'pdf_to_word': 'pdf-to-word.html',
    'jpg_to_pdf': 'jpg-to-pdf.html',
    'pdf_to_jpg': 'pdf-to-jpg.html',
    'rotate_pdf': 'rotate.html',
    'protect-pdf': 'protect.html',
    'unlock_pdf': 'unlock.html',
    'pdf_add_watermark': 'watermark.html',
    'add_pdf_page_number': 'page-numbers.html',
    'organize-pdf': 'organize.html',
    'html-to-pdf': 'html-to-pdf.html',
    'extract-text': 'extract-text.html',
    'remove-pages': 'delete-pages.html',
    'compare-pdf': 'compare.html',
    'powerpoint_to_pdf': 'powerpoint-to-pdf.html',
    'pdf_to_powerpoint': 'pdf-to-powerpoint.html',
    'excel_to_pdf': 'excel-to-pdf.html',
    'pdf_to_excel': 'pdf-to-excel.html',
    'convert-pdf-to-pdfa': 'pdf-to-pdfa.html',
    'repair-pdf': 'repair-pdf.html',
    'ocr-pdf': 'ocr-pdf.html',
    'pdf-summarize': 'pdf-summarize.html',
    'translate-pdf': 'translate-pdf.html',
    'sign-pdf': 'sign-pdf.html'
}

TRANSLATIONS = {
    'es': {
        'Every tool you need to work with PDFs': 'Todas las herramientas PDF que necesitas en un solo lugar',
        'Every tool you need to use PDFs, at your fingertips. All are 100% FREE and easy to use!': 'Todas las herramientas PDF que necesitas al alcance de tu mano. 100% GRATIS y fáciles de usar.',
        'Merge, split, compress, convert, rotate, protect, unlock and watermark PDFs with just a few clicks.': 'Une, divide, comprime, convierte, rota, protege, desbloquea y añade marcas de agua con unos pocos clics.',
        'Merge PDF': 'Unir PDF',
        'Split PDF': 'Dividir PDF',
        'Compress PDF': 'Comprimir PDF',
        'Word to PDF': 'Word a PDF',
        'PDF to Word': 'PDF a Word',
        'JPG to PDF': 'JPG a PDF',
        'PDF to JPG': 'PDF a JPG',
        'Rotate PDF': 'Rotar PDF',
        'Protect PDF': 'Proteger PDF',
        'Unlock PDF': 'Desbloquear PDF',
        'Watermark': 'Marca de agua',
        'Add Watermark': 'Añadir Marca de Agua',
        'Page Numbers': 'Números de página',
        'Add page numbers': 'Añadir números de página',
        'Organize PDF': 'Organizar PDF',
        'HTML to PDF': 'HTML a PDF',
        'Extract Text': 'Extraer texto',
        'Delete Pages': 'Eliminar páginas',
        'Compare PDF': 'Comparar PDF',
        'PowerPoint to PDF': 'PowerPoint a PDF',
        'PDF to PowerPoint': 'PDF a PowerPoint',
        'Excel to PDF': 'Excel a PDF',
        'PDF to Excel': 'PDF a Excel',
        'PDF to PDF/A': 'PDF a PDF/A',
        'Repair PDF': 'Reparar PDF',
        'OCR PDF': 'OCR PDF',
        'AI Summarizer': 'Resumidor de IA',
        'Translate PDF': 'Traducir PDF',
        'Sign PDF': 'Firmar PDF',
        'All PDF tools': 'Todas las herramientas PDF',
        'Convert PDF': 'Convertir PDF',
        'Log in': 'Iniciar sesión',
        'Login': 'Acceso',
        'Sign up': 'Registrarse',
        'Home': 'Inicio',
        'Combine PDF files in the order you want with the easiest PDF merger available.': 'Combina archivos PDF en el orden que desees con la fusionadora más fácil.',
        'Extract page ranges or separate all PDF pages to save as multiple PDF files.': 'Extrae rangos de páginas o separa todas las páginas PDF en archivos individuales.',
        'Reduce file size of your PDF while keeping the best quality and formatting.': 'Reduce el tamaño de tu PDF manteniendo la mejor calidad y formato.',
        'Convert Microsoft Word DOCX documents into clean PDF files.': 'Convierte tus documentos DOCX a archivos PDF limpios.',
        'Convert PDF files to editable Microsoft Word DOCX documents easily.': 'Convierte PDFs a documentos Word editables fácilmente.',
        'Convert JPG, PNG, WEBP and BMP images to PDF quickly.': 'Convierte imágenes JPG, PNG, WEBP y BMP a PDF de forma rápida.',
        'Extract all embedded images or convert each PDF page into a high-quality JPEG.': 'Extrae imágenes incrustadas o convierte cada página en un JPEG de alta calidad.',
        'Rotate your PDF pages to match portrait or landscape formats.': 'Rota tus páginas PDF a formato horizontal o vertical.',
        'Encrypt your PDF with a strong security password to restrict access.': 'Protege tu PDF con una contraseña de seguridad fuerte.',
        'Remove password protection from secured PDFs so you can open them anytime.': 'Quita contraseñas de PDFs para abrirlos libremente.',
        'Embed text overlays with custom opacity, angles and formats over PDF pages.': 'Añade textos superpuestos con opacidad y ángulos personalizados.',
        'Add page numbers to your document layout in custom styles and locations.': 'Agrega números de página con estilos y posiciones a elegir.',
        'Rearrange, rotate, or delete specific pages from your document in seconds.': 'Reordena, rota o elimina páginas del documento en segundos.',
        'Convert custom HTML code or text fragments into standard PDF documents.': 'Convierte fragmentos HTML a documentos PDF estándar.',
        'Read and extract all raw text strings from PDF pages into a TXT file.': 'Extrae todo el texto plano de tus PDFs a un archivo de texto.',
        'Remove unwanted pages or page ranges from a PDF document.': 'Elimina páginas o rangos no deseados de un documento PDF.',
        'Other products': 'Otros productos',
        'Solutions': 'Soluciones',
        'Applications': 'Aplicaciones',
        'Pricing': 'Precios',
        'Security': 'Seguridad',
        'Features': 'Funciones',
        'About us': 'Sobre nosotros',
        'Help': 'Ayuda',
        'Language': 'Idioma',
        'Free online service to work with PDF files completely free and easy to use.': 'Servicio online gratuito y fácil de usar para trabajar con tus archivos PDF.',
        'We encrypt all transactions using 256-bit SSL technology.': 'Ciframos todas las transacciones usando tecnologías SSL de 256 bits.',
        'Your files are automatically purged from our systems after 15 minutes of inactivity.': 'Tus archivos son eliminados automáticamente tras 15 minutos.'
    },
    'fr': {
        'Every tool you need to work with PDFs': 'Tous les outils PDF dont vous avez besoin au même endroit',
        'Every tool you need to use PDFs, at your fingertips. All are 100% FREE and easy to use!': 'Tous les outils PDF au bout de vos doigts. 100% GRATUIT et facile d\'utilisation !',
        'Merge, split, compress, convert, rotate, protect, unlock and watermark PDFs with just a few clicks.': 'Fusionnez, divisez, compressez, convertissez, pivotez, protégez et déverrouillez des PDFs en quelques clics.',
        'Merge PDF': 'Fusionner PDF',
        'Split PDF': 'Diviser PDF',
        'Compress PDF': 'Compresser PDF',
        'Word to PDF': 'Word en PDF',
        'PDF to Word': 'PDF en Word',
        'JPG to PDF': 'JPG en PDF',
        'PDF to JPG': 'PDF en JPG',
        'Rotate PDF': 'Pivoter PDF',
        'Protect PDF': 'Protéger PDF',
        'Unlock PDF': 'Déverrouiller PDF',
        'Watermark': 'Filigrane',
        'Add Watermark': 'Ajouter filigrane',
        'Page Numbers': 'Numérotation',
        'Add page numbers': 'Ajouter numéros de page',
        'Organize PDF': 'Organiser PDF',
        'HTML to PDF': 'HTML en PDF',
        'Extract Text': 'Extraire texte',
        'Delete Pages': 'Supprimer pages',
        'Compare PDF': 'Comparer PDF',
        'PowerPoint to PDF': 'PowerPoint en PDF',
        'PDF to PowerPoint': 'PDF en PowerPoint',
        'Excel to PDF': 'Excel en PDF',
        'PDF to Excel': 'PDF en Excel',
        'PDF to PDF/A': 'PDF en PDF/A',
        'Repair PDF': 'Réparer PDF',
        'OCR PDF': 'OCR PDF',
        'AI Summarizer': 'Résumeur IA',
        'Translate PDF': 'Traduire PDF',
        'Sign PDF': 'Signer PDF',
        'All PDF tools': 'Tous les outils PDF',
        'Convert PDF': 'Convertir PDF',
        'Log in': 'Connexion',
        'Login': 'Accès',
        'Sign up': 'S\'inscrire',
        'Home': 'Accueil'
    },
    'de': {
        'Every tool you need to work with PDFs': 'Jedes PDF-Werkzeug, das Sie benötigen, an einem Ort',
        'Every tool you need to use PDFs, at your fingertips. All are 100% FREE and easy to use!': 'Alle PDF-Werkzeuge für Sie bereit. 100% KOSTENLOS und einfach zu bedienen!',
        'Merge PDF': 'PDF zusammenfügen',
        'Split PDF': 'PDF teilen',
        'Compress PDF': 'PDF komprimieren',
        'Word to PDF': 'Word in PDF',
        'PDF to Word': 'PDF in Word',
        'JPG to PDF': 'JPG in PDF',
        'PDF to JPG': 'PDF in JPG',
        'Rotate PDF': 'PDF drehen',
        'Protect PDF': 'PDF schützen',
        'Unlock PDF': 'PDF entsperren',
        'Watermark': 'Wasserzeichen',
        'Add Watermark': 'Wasserzeichen hinzufügen',
        'Page Numbers': 'Seitenzahlen',
        'Add page numbers': 'Seitenzahlen hinzufügen',
        'Organize PDF': 'PDF organisieren',
        'HTML to PDF': 'HTML in PDF',
        'Extract Text': 'Text extrahieren',
        'Delete Pages': 'Seiten löschen',
        'Compare PDF': 'PDF vergleichen',
        'PowerPoint to PDF': 'PowerPoint in PDF',
        'PDF to PowerPoint': 'PDF in PowerPoint',
        'Excel to PDF': 'Excel in PDF',
        'PDF to Excel': 'PDF in Excel',
        'PDF to PDF/A': 'PDF in PDF/A',
        'Repair PDF': 'PDF reparieren',
        'OCR PDF': 'OCR PDF',
        'AI Summarizer': 'KI-Zusammenfassung',
        'Translate PDF': 'PDF übersetzen',
        'Sign PDF': 'PDF unterschreiben',
        'All PDF tools': 'Alle PDF-Werkzeuge',
        'Log in': 'Einloggen',
        'Sign up': 'Registrieren',
        'Home': 'Startseite'
    },
    'pt': {
        'Every tool you need to work with PDFs': 'Todas as ferramentas PDF necessárias em um só lugar',
        'Merge PDF': 'Mesclar PDF',
        'Split PDF': 'Dividir PDF',
        'Compress PDF': 'Comprimir PDF',
        'Word to PDF': 'Word para PDF',
        'PDF to Word': 'PDF para Word',
        'JPG to PDF': 'JPG para PDF',
        'PDF to JPG': 'PDF para JPG',
        'Rotate PDF': 'Rotacionar PDF',
        'Protect PDF': 'Proteger PDF',
        'Unlock PDF': 'Desbloquear PDF',
        'Watermark': 'Marca d\'água',
        'Page Numbers': 'Números de página',
        'Organize PDF': 'Organizar PDF',
        'HTML to PDF': 'HTML para PDF',
        'Extract Text': 'Extrair texto',
        'Delete Pages': 'Excluir páginas',
        'Compare PDF': 'Comparar PDF',
        'PowerPoint to PDF': 'PowerPoint para PDF',
        'PDF to PowerPoint': 'PDF para PowerPoint',
        'Excel to PDF': 'Excel para PDF',
        'PDF to Excel': 'PDF para Excel',
        'PDF to PDF/A': 'PDF para PDF/A',
        'Repair PDF': 'Reparar PDF',
        'OCR PDF': 'OCR PDF',
        'AI Summarizer': 'Resumidor de IA',
        'Translate PDF': 'Traduzir PDF',
        'Sign PDF': 'Assinar PDF',
        'All PDF tools': 'Todas as ferramentas PDF',
        'Log in': 'Entrar',
        'Sign up': 'Cadastrar',
        'Home': 'Início'
    },
    'hi': {
        'Every tool you need to work with PDFs': 'पीडीएफ पर काम करने के लिए सभी आवश्यक उपकरण एक जगह',
        'Merge PDF': 'पीडीएफ मर्ज करें',
        'Split PDF': 'पीडीएफ विभाजित करें',
        'Compress PDF': 'पीडीएफ कंप्रेस करें',
        'Word to PDF': 'वर्ड से पीडीएफ',
        'PDF to Word': 'पीडीएफ से वर्ड',
        'JPG to PDF': 'जेपीजी से पीडीएफ',
        'PDF to JPG': 'पीडीएफ से जेपीजी',
        'Rotate PDF': 'पीडीएफ घुमाएँ',
        'Protect PDF': 'पीडीएफ सुरक्षित करें',
        'Unlock PDF': 'पीडीएफ अनलॉक करें',
        'Watermark': 'वॉटरमार्क',
        'Page Numbers': 'पेज नंबर',
        'Organize PDF': 'पीडीएफ व्यवस्थित करें',
        'HTML to PDF': 'एचटीएमएल से पीडीएफ',
        'Extract Text': 'टेक्स्ट निकालें',
        'Delete Pages': 'पेज हटाएं',
        'Compare PDF': 'पीडीएफ तुलना करें',
        'PowerPoint to PDF': 'पॉवरपॉइंट से पीडीएफ',
        'PDF to PowerPoint': 'पीडीएफ से पॉवरपॉइंट',
        'Excel to PDF': 'एक्सेल से पीडीएफ',
        'PDF to Excel': 'पीडीएफ से एक्सेल',
        'PDF to PDF/A': 'पीडीएफ से पीडीएफ/ए',
        'Repair PDF': 'पीडीएफ रिपेयर करें',
        'OCR PDF': 'ओसीआर पीडीएफ',
        'AI Summarizer': 'एिआई सारांश',
        'Translate PDF': 'अनुवाद करें',
        'Sign PDF': 'हस्ताक्षर करें',
        'All PDF tools': 'सभी पीडीएफ टूल्स',
        'Log in': 'लॉगिन',
        'Sign up': 'साइन अप',
        'Home': 'होम'
    }
}
# Serve explicit static assets from root
@app.route('/style.css')
def serve_css():
    return send_from_directory(BASE_DIR, 'style.css')

@app.route('/ad-manager.js')
def serve_ad_manager():
    return send_from_directory(BASE_DIR, 'ad-manager.js')

# Serve pages cleanly matching URLs and localized parameters
@app.route('/')
@app.route('/<lang>')
def index_route(lang=None):
    if lang and lang not in SUPPORTED_LANGS:
        if lang in SLUG_TO_FILE:
            return serve_tool_page(lang, None)
        return "Page not found.", 404
    return serve_tool_page('', lang)

@app.route('/<lang>/<slug>')
def lang_tool_route(lang, slug):
    if lang in SUPPORTED_LANGS and slug in SLUG_TO_FILE:
        return serve_tool_page(slug, lang)
    elif lang in SLUG_TO_FILE:
        return serve_tool_page(lang, None)
    return "Page not found.", 404

def serve_tool_page(slug, lang=None):
    # Backward compatibility redirects
    legacy_redirects = {
        'merge': 'merge_pdf',
        'split': 'split_pdf',
        'compress': 'compress_pdf',
        'word-to-pdf': 'word_to_pdf',
        'pdf-to-word': 'pdf_to_word',
        'jpg-to-pdf': 'jpg_to_pdf',
        'pdf-to-jpg': 'pdf_to_jpg',
        'rotate': 'rotate_pdf',
        'protect': 'protect-pdf',
        'unlock': 'unlock_pdf',
        'watermark': 'pdf_add_watermark',
        'page-numbers': 'add_pdf_page_number',
        'organize': 'organize-pdf',
        'delete-pages': 'remove-pages',
        'compare': 'compare-pdf'
    }
    if slug in legacy_redirects:
        target = legacy_redirects[slug]
        url = f"/{lang}/{target}" if lang else f"/{target}"
        return redirect(url, code=301)

    filename = SLUG_TO_FILE.get(slug)
    if not filename:
        # Fallback check if first parameter was actually slug without lang
        if slug in SUPPORTED_LANGS:
            filename = 'index.html'
        else:
            return "Page not found.", 404
            
    filepath = os.path.join(BASE_DIR, filename)
    if not os.path.exists(filepath):
        return "File not found.", 404
        
    with open(filepath, 'r', encoding='utf-8') as f:
        html = f.read()
        
    # --- Inject Canonical & Hreflang SEO Metadata ---
    root_domain = "https://ilovespdfs.in"
    lang_sub = f"{lang}/" if lang and lang != 'en' else ""
    canonical_url = f"{root_domain}/{lang_sub}{slug}"
    
    hreflangs = []
    hreflangs.append(f'<link rel="alternate" hreflang="x-default" href="{root_domain}/{slug}" />')
    hreflangs.append(f'<link rel="alternate" hreflang="en" href="{root_domain}/{slug}" />')
    for l in SUPPORTED_LANGS:
        if l != 'en':
            hreflangs.append(f'<link rel="alternate" hreflang="{l}" href="{root_domain}/{l}/{slug}" />')
            
    hreflang_tags = "\n    ".join(hreflangs)
    
    # Strip existing canonical and inject new
    import re
    html = re.sub(r'<link\s+rel=["\']canonical["\'].*?>', '', html, flags=re.IGNORECASE)
    injection = f'<link rel="canonical" href="{canonical_url}" />\n    {hreflang_tags}'
    html = re.sub(r'(<head\b[^>]*>)', r'\1\n    ' + injection, html, flags=re.IGNORECASE)
    
    # --- Server-Side Dictionary Translation ---
    if lang and lang in TRANSLATIONS and lang != 'en':
        trans_dict = TRANSLATIONS[lang]
        sorted_keys = sorted(trans_dict.keys(), key=len, reverse=True)
        for key in sorted_keys:
            html = html.replace(key, trans_dict[key])
            
    return html

# Compatibility fallbacks
@app.route('/merge')
@app.route('/split')
@app.route('/compress')
@app.route('/word-to-pdf')
@app.route('/pdf-to-word')
@app.route('/jpg-to-pdf')
@app.route('/pdf-to-jpg')
@app.route('/rotate')
@app.route('/protect')
@app.route('/unlock')
@app.route('/watermark')
@app.route('/page-numbers')
@app.route('/organize')
@app.route('/html-to-pdf')
@app.route('/extract-text')
@app.route('/delete-pages')
@app.route('/compare')
def legacy_page_redirect():
    slug = request.path.strip('/')
    return serve_tool_page(slug, None)

@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({'error': 'File is too large. Maximum size is 50MB.'}), 413

# --- PDF Tools API ---

@app.route('/api/merge', methods=['POST'])
def merge_pdfs():
    try:
        saved_paths = process_upload(request.files.getlist('files'), ['.pdf'], max_files=100)
        if len(saved_paths) < 2:
            raise ValueError('At least 2 PDF files are required for merging.')
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    try:

        # Merge
        merger = PdfWriter()
        for path in saved_paths:
            merger.append(path)

        out_filename = f"merged_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        with open(out_path, 'wb') as f_out:
            merger.write(f_out)
        merger.close()

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        # Clean up input uploads immediately
        for path in saved_paths:
            try:
                os.remove(path)
            except:
                pass

@app.route('/api/split', methods=['POST'])
def split_pdf():
    file = request.files.get('file')
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400

    ranges_str = request.form.get('ranges', '1-end') # e.g. "1-3, 4-5" or "1-end"
    
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)

    try:
        reader = PdfReader(input_path)
        total_pages = len(reader.pages)

        # Parse ranges
        ranges = []
        parts = [p.strip() for p in ranges_str.split(',') if p.strip()]
        
        try:
            for part in parts:
                if '-' in part:
                    start_s, end_s = part.split('-', 1)
                    start_s = start_s.strip()
                    end_s = end_s.strip()
                    
                    start = int(start_s) if start_s else 1
                    if end_s.lower() == 'end' or not end_s:
                        end = total_pages
                    else:
                        end = int(end_s)
                else:
                    start = int(part)
                    end = start
                
                # Bound checking (1-based index)
                start = max(1, min(start, total_pages))
                end = max(1, min(end, total_pages))
                if start <= end:
                    ranges.append((start, end))
                else:
                    ranges.append((end, start))
        except ValueError:
            return jsonify({'error': 'Invalid range format. Please use integers like "1-3" or "1-end".'}), 400

        if not ranges:
            return jsonify({'error': 'Invalid ranges specified.'}), 400

        # If only one range, return single PDF
        if len(ranges) == 1:
            start, end = ranges[0]
            writer = PdfWriter()
            for p_num in range(start - 1, end):
                writer.add_page(reader.pages[p_num])
            
            out_filename = f"split_{start}-{end}_{uuid.uuid4().hex[:8]}.pdf"
            out_path = os.path.join(PROCESSED_FOLDER, out_filename)
            with open(out_path, 'wb') as f_out:
                writer.write(f_out)
            
            return jsonify({'download_url': f'/download/{out_filename}'})
        else:
            # If multiple ranges, create a ZIP
            zip_filename = f"split_{uuid.uuid4().hex[:8]}.zip"
            zip_path = os.path.join(PROCESSED_FOLDER, zip_filename)
            
            with zipfile.ZipFile(zip_path, 'w') as zf:
                for idx, (start, end) in enumerate(ranges):
                    writer = PdfWriter()
                    for p_num in range(start - 1, end):
                        writer.add_page(reader.pages[p_num])
                    
                    pdf_bytes = io.BytesIO()
                    writer.write(pdf_bytes)
                    pdf_bytes.seek(0)
                    zf.writestr(f"split_range_{idx+1}_{start}-{end}.pdf", pdf_bytes.read())

            return jsonify({'download_url': f'/download/{zip_filename}'})

    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/compress', methods=['POST'])
def compress_pdf():
    try:
        saved_paths = process_upload(request.files.get('file'), ['.pdf'])
        input_path = saved_paths[0]
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    try:
        # Try PyMuPDF optimization first (highly effective)
        try:
            doc = fitz.open(input_path)
            out_filename = f"compressed_{uuid.uuid4().hex[:8]}.pdf"
            out_path = os.path.join(PROCESSED_FOLDER, out_filename)
            doc.save(out_path, garbage=4, deflate=True, clean=True)
            doc.close()
        except Exception:
            # Fallback to pypdf stream compression
            reader = PdfReader(input_path)
            writer = PdfWriter()
            for page in reader.pages:
                added_page = writer.add_page(page)
                added_page.compress_content_streams()
            out_filename = f"compressed_{uuid.uuid4().hex[:8]}.pdf"
            out_path = os.path.join(PROCESSED_FOLDER, out_filename)
            with open(out_path, 'wb') as f_out:
                writer.write(f_out)

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/jpg_to_pdf', methods=['POST'])
def jpg_to_pdf():
    try:
        saved_paths = process_upload(request.files.getlist('files'), ['.jpg', '.jpeg', '.png', '.webp', '.bmp'], max_files=100)
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    try:

        # Convert images to PDF
        images = []
        for path in saved_paths:
            img = Image.open(path)
            if img.mode != 'RGB':
                img = img.convert('RGB')
            images.append(img)

        out_filename = f"converted_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        
        if images:
            images[0].save(out_path, save_all=True, append_images=images[1:])

        # Close all image handlers
        for img in images:
            img.close()

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        for path in saved_paths:
            try:
                os.remove(path)
            except:
                pass

@app.route('/api/pdf_to_jpg', methods=['POST'])
def pdf_to_jpg():
    try:
        saved_paths = process_upload(request.files.get('file'), ['.pdf'])
        input_path = saved_paths[0]
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    try:
        doc = fitz.open(input_path)
        zip_filename = f"images_{uuid.uuid4().hex[:8]}.zip"
        zip_path = os.path.join(PROCESSED_FOLDER, zip_filename)

        with zipfile.ZipFile(zip_path, 'w') as zf:
            for idx, page in enumerate(doc):
                # Render page to image (pixmap)
                pix = page.get_pixmap(dpi=150)
                # Convert colorspace to RGB if it is not RGB
                if pix.colorspace and pix.colorspace.n != 3:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                img_data = pix.tobytes("jpg")
                zf.writestr(f"page_{idx+1}.jpg", img_data)

        doc.close()
        return jsonify({'download_url': f'/download/{zip_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/rotate', methods=['POST'])
def rotate_pdf():
    file = request.files.get('file')
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400

    try:
        angle = int(request.form.get('angle', 90))
        if angle not in [90, 180, 270, 360, -90, -180, -270]:
            angle = 90
    except ValueError:
        angle = 90

    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)

    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()

        for page in reader.pages:
            page.rotate(angle)
            writer.add_page(page)

        out_filename = f"rotated_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        with open(out_path, 'wb') as f_out:
            writer.write(f_out)

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/protect', methods=['POST'])
def protect_pdf():
    file = request.files.get('file')
    password = request.form.get('password')
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    if not password:
        return jsonify({'error': 'Password is required to protect the PDF.'}), 400

    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)

    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()
        writer.append_pages_from_reader(reader)
        writer.encrypt(password)

        out_filename = f"protected_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        with open(out_path, 'wb') as f_out:
            writer.write(f_out)

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/unlock', methods=['POST'])
def unlock_pdf():
    file = request.files.get('file')
    password = request.form.get('password', '')
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400

    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)

    try:
        reader = PdfReader(input_path)
        
        if reader.is_encrypted:
            success = reader.decrypt(password)
            if not success:
                return jsonify({'error': 'Incorrect password. Could not decrypt PDF.'}), 400
                
        writer = PdfWriter()
        writer.append_pages_from_reader(reader)

        out_filename = f"unlocked_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        with open(out_path, 'wb') as f_out:
            writer.write(f_out)

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/watermark', methods=['POST'])
def watermark_pdf():
    file = request.files.get('file')
    text = request.form.get('text', 'i loves pdf')
    try:
        opacity = float(request.form.get('opacity', 0.3))
        opacity = max(0.0, min(opacity, 1.0))
    except ValueError:
        opacity = 0.3
        
    try:
        rotation = int(request.form.get('rotation', 45))
    except ValueError:
        rotation = 45

    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400

    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)

    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()

        for page in reader.pages:
            box = page.cropbox
            width = float(box.width)
            height = float(box.height)
            
            # Generate watermark overlay PDF in memory matching page dimensions
            packet = io.BytesIO()
            can = canvas.Canvas(packet, pagesize=(width, height))
            can.setFont("Helvetica", 40)
            can.setFillColorRGB(0.7, 0.7, 0.7, alpha=opacity)
            can.saveState()
            
            # Center in page dynamically
            can.translate(width / 2, height / 2)
            can.rotate(rotation)
            can.drawCentredString(0, 0, text)
            can.restoreState()
            can.save()
            packet.seek(0)
            
            watermark_reader = PdfReader(packet)
            watermark_page = watermark_reader.pages[0]
            
            page.merge_page(watermark_page)
            writer.add_page(page)

        out_filename = f"watermarked_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        with open(out_path, 'wb') as f_out:
            writer.write(f_out)

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/page_numbers', methods=['POST'])
def add_page_numbers():
    file = request.files.get('file')
    position = request.form.get('position', 'bottom-right') # bottom-left, bottom-center, bottom-right, top-left, top-center, top-right
    try:
        start_num = int(request.form.get('start_number', 1))
    except ValueError:
        start_num = 1
    
    label_format = request.form.get('format', 'Page {n}') # e.g. "Page {n}" or "{n}"

    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400

    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)

    try:
        reader = PdfReader(input_path)
        total_pages = len(reader.pages)
        writer = PdfWriter()

        for idx in range(total_pages):
            page = reader.pages[idx]
            
            # Check cropbox size of the page
            box = page.cropbox
            width = float(box.width)
            height = float(box.height)
            
            # Generate overlay page number PDF
            packet = io.BytesIO()
            can = canvas.Canvas(packet, pagesize=(width, height))
            can.setFont("Helvetica", 10)
            can.setFillColorRGB(0.3, 0.3, 0.3)
            
            num_text = label_format.replace('{n}', str(start_num + idx))
            
            # Determine x, y based on margin and size
            margin = 36
            if 'top' in position:
                y = height - margin
            else:
                y = margin
                
            if 'left' in position:
                x = margin
                can.drawString(x, y, num_text)
            elif 'center' in position:
                x = width / 2
                can.drawCentredString(x, y, num_text)
            else: # right
                x = width - margin
                can.drawRightString(x, y, num_text)
                
            can.save()
            packet.seek(0)
            
            number_pdf = PdfReader(packet)
            page.merge_page(number_pdf.pages[0])
            writer.add_page(page)

        out_filename = f"numbered_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        with open(out_path, 'wb') as f_out:
            writer.write(f_out)

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/organize', methods=['POST'])
def organize_pdf():
    file = request.files.get('file')
    order_str = request.form.get('order')  # comma-separated 0-based index list e.g. "0,2,1"
    rotations_str = request.form.get('rotations') # comma-separated angles e.g. "90,0,180"
    
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    if not order_str:
        return jsonify({'error': 'Page ordering config is required.'}), 400

    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)

    try:
        reader = PdfReader(input_path)
        total_pages = len(reader.pages)
        
        try:
            indices = [int(i.strip()) for i in order_str.split(',') if i.strip()]
        except ValueError:
            return jsonify({'error': 'Invalid page ordering format.'}), 400

        rotations = []
        if rotations_str:
            try:
                rotations = [int(r.strip()) for r in rotations_str.split(',') if r.strip()]
            except ValueError:
                pass

        writer = PdfWriter()
        for idx_in_list, idx in enumerate(indices):
            if 0 <= idx < total_pages:
                page = writer.add_page(reader.pages[idx])
                if rotations and idx_in_list < len(rotations):
                    angle = rotations[idx_in_list]
                    if angle in [90, 180, 270, 360, -90, -180, -270]:
                        page.rotate(angle)

        out_filename = f"organized_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        with open(out_path, 'wb') as f_out:
            writer.write(f_out)

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/word_to_pdf', methods=['POST'])
def word_to_pdf():
    try:
        saved_paths = process_upload(request.files.get('file'), ['.doc', '.docx'])
        input_path = saved_paths[0]
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    try:
        # Load the Word document
        doc = Document(input_path)
        
        # Build ReportLab elements
        out_filename = f"converted_word_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        
        pdf_doc = SimpleDocTemplate(out_path, pagesize=letter,
                                    rightMargin=72, leftMargin=72,
                                    topMargin=72, bottomMargin=72)
        
        styles = getSampleStyleSheet()
        
        # Create a clean normal style
        normal_style = ParagraphStyle(
            name='WordNormal',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=11,
            leading=14,
            spaceAfter=8
        )
        
        heading1_style = ParagraphStyle(
            name='WordHeading1',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=18,
            leading=22,
            spaceAfter=12,
            spaceBefore=12
        )
        
        story = []
        
        from reportlab.platypus import Table as RLTable, TableStyle as RLTableStyle
        from reportlab.lib import colors
        
        for item in iter_block_items(doc):
            if isinstance(item, DocxParagraph):
                text = item.text.strip()
                if not text:
                    story.append(Spacer(1, 10))
                    continue
                
                style = normal_style
                if item.style and item.style.name:
                    if item.style.name.startswith('Heading 1'):
                        style = heading1_style
                    elif item.style.name.startswith('Heading'):
                        style = ParagraphStyle(
                            name='WordHeadingSub',
                            parent=styles['Heading2'],
                            fontName='Helvetica-Bold',
                            fontSize=14,
                            leading=18,
                            spaceAfter=10,
                            spaceBefore=10
                        )
                
                safe_text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe_text, style))
            elif isinstance(item, DocxTable):
                table_data = []
                for row in item.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = "\n".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                        safe_cell_text = cell_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        row_data.append(Paragraph(safe_cell_text or ' ', normal_style))
                    table_data.append(row_data)
                
                if table_data:
                    rl_table = RLTable(table_data, colWidths=None)
                    rl_table.setStyle(RLTableStyle([
                        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
                        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                        ('VALIGN', (0,0), (-1,-1), 'TOP'),
                        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                        ('TOPPADDING', (0,0), (-1,-1), 6),
                        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
                        ('LEFTPADDING', (0,0), (-1,-1), 6),
                        ('RIGHTPADDING', (0,0), (-1,-1), 6),
                    ]))
                    story.append(rl_table)
                    story.append(Spacer(1, 12))
            
        # Build PDF file
        pdf_doc.build(story)

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

# --- New PDF Tools API ---

@app.route('/api/pdf_to_word', methods=['POST'])
def pdf_to_word():
    try:
        saved_paths = process_upload(request.files.get('file'), ['.pdf'])
        input_path = saved_paths[0]
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    try:
        doc = fitz.open(input_path)
        word_doc = Document()
        
        for page in doc:
            blocks = page.get_text("blocks")
            blocks.sort(key=lambda b: (b[1], b[0]))
            for block in blocks:
                block_text = block[4].strip()
                if block_text:
                    paragraph_text = " ".join([line.strip() for line in block_text.split("\n") if line.strip()])
                    word_doc.add_paragraph(paragraph_text)
        
        out_filename = f"converted_pdf_{uuid.uuid4().hex[:8]}.docx"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        word_doc.save(out_path)
        doc.close()
        
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/html_to_pdf', methods=['POST'])
def html_to_pdf():
    file = request.files.get('file')
    html_code = request.form.get('html_code', '')

    html_content = ""
    if file and file.filename != '':
        # Read uploaded file
        try:
            html_content = file.read().decode('utf-8')
        except Exception as e:
            return jsonify({'error': f'Failed to read HTML file: {str(e)}'}), 400
    else:
        html_content = html_code

    if not html_content.strip():
        return jsonify({'error': 'HTML content or file is required.'}), 400

    out_filename = f"html_pdf_{uuid.uuid4().hex[:8]}.pdf"
    out_path = os.path.join(PROCESSED_FOLDER, out_filename)

    try:
        # Attempt conversion via fitz HTML support
        doc = fitz.open("html", html_content)
        pdf_bytes = doc.convert_to_pdf()
        with open(out_path, 'wb') as f_out:
            f_out.write(pdf_bytes)
        doc.close()
    except Exception as fitz_err:
        # Fallback to reportlab custom parser
        try:
            pdf_doc = SimpleDocTemplate(out_path, pagesize=letter)
            styles = getSampleStyleSheet()
            normal_style = ParagraphStyle(
                name='HTMLNormal',
                parent=styles['Normal'],
                fontName='Helvetica',
                fontSize=11,
                leading=14,
                spaceAfter=8
            )
            parser = ReportLabHTMLParser(styles, normal_style)
            parser.feed(html_content)
            parser.flush()
            if not parser.story:
                parser.story.append(Paragraph("Empty Document", normal_style))
            pdf_doc.build(parser.story)
        except Exception as rl_err:
            return jsonify({'error': f'PDF conversion failed: {str(fitz_err)} | Fallback: {str(rl_err)}'}), 500

    return jsonify({'download_url': f'/download/{out_filename}'})

@app.route('/api/extract_text', methods=['POST'])
def extract_text():
    try:
        saved_paths = process_upload(request.files.get('file'), ['.pdf'])
        input_path = saved_paths[0]
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    try:
        doc = fitz.open(input_path)
        text_content = ""
        for page in doc:
            text_content += page.get_text() + "\n"
        
        out_filename = f"extracted_{uuid.uuid4().hex[:8]}.txt"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        with open(out_path, 'w', encoding='utf-8') as f_out:
            f_out.write(text_content)
        doc.close()
        
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/delete_pages', methods=['POST'])
def delete_pdf_pages():
    file = request.files.get('file')
    pages_str = request.form.get('pages', '')
    
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    if not pages_str:
        return jsonify({'error': 'Pages to delete configuration is required.'}), 400

    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)

    try:
        reader = PdfReader(input_path)
        total_pages = len(reader.pages)
        
        # Parse ranges to delete
        deleted_indices = set()
        parts = [p.strip() for p in pages_str.split(',') if p.strip()]
        for part in parts:
            if '-' in part:
                start_s, end_s = part.split('-', 1)
                start = int(start_s) if start_s.strip() else 1
                if end_s.lower() == 'end' or not end_s.strip():
                    end = total_pages
                else:
                    end = int(end_s)
                for p in range(min(start, end), max(start, end) + 1):
                    if 1 <= p <= total_pages:
                        deleted_indices.add(p - 1)
            else:
                p = int(part)
                if 1 <= p <= total_pages:
                    deleted_indices.add(p - 1)
        
        if len(deleted_indices) >= total_pages:
            return jsonify({'error': 'Cannot delete all pages from the PDF.'}), 400
            
        writer = PdfWriter()
        for idx in range(total_pages):
            if idx not in deleted_indices:
                writer.add_page(reader.pages[idx])
                
        out_filename = f"deleted_pages_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        with open(out_path, 'wb') as f_out:
            writer.write(f_out)
            
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

# --- File Serving and Downloads ---

@app.route('/download/<filename>')
def download_file(filename):
    filename = os.path.basename(filename)
    file_path = os.path.join(PROCESSED_FOLDER, filename)
    if os.path.exists(file_path):
        return send_from_directory(PROCESSED_FOLDER, filename, as_attachment=True)
    else:
        return "File not found or expired.", 404

@app.route('/compare-pdf')
def compare_pdf_redirect():
    return serve_tool_page('compare-pdf', None)

# --- Technical SEO and Crawl Control ---

@app.route('/robots.txt')
def serve_robots():
    return send_from_directory(BASE_DIR, 'robots.txt', mimetype='text/plain')

@app.route('/sitemap.xml')
def serve_sitemap():
    root = "https://ilovespdfs.in"
    xml = ['<?xml version="1.0" encoding="UTF-8"?>',
           '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">']
    
    for slug in SLUG_TO_FILE:
        path = f"/{slug}" if slug else ""
        xml.append('  <url>')
        xml.append(f'    <loc>{root}{path}</loc>')
        xml.append('    <changefreq>daily</changefreq>')
        xml.append('    <priority>0.9</priority>')
        xml.append('  </url>')
        
        for lang in SUPPORTED_LANGS:
            if lang != 'en':
                path_lang = f"/{lang}/{slug}" if slug else f"/{lang}"
                xml.append('  <url>')
                xml.append(f'    <loc>{root}{path_lang}</loc>')
                xml.append('    <changefreq>daily</changefreq>')
                xml.append('    <priority>0.8</priority>')
                xml.append('  </url>')
                
    xml.append('</urlset>')
    return "\n".join(xml), 200, {'Content-Type': 'application/xml'}

@app.after_request
def add_header(response):
    path = request.path
    if (path.endswith('.css') or path.endswith('.js') or 
        path.endswith('.png') or path.endswith('.jpg') or 
        path.endswith('.svg') or path.endswith('.webp') or 
        path.endswith('.ico') or path.endswith('.woff2')):
        response.headers['Cache-Control'] = 'public, max-age=31536000, immutable'
    else:
        response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    return response

@app.route('/api/powerpoint_to_pdf', methods=['POST'])
def powerpoint_to_pdf():
    file = request.files.get('file')
    if not file or file.filename == '':
        return jsonify({'error': 'PowerPoint file is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pptx'))
    file.save(input_path)
    try:
        out_filename = f"converted_pptx_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        convert_pptx_to_pdf(input_path, out_path)
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/pdf_to_powerpoint', methods=['POST'])
def pdf_to_powerpoint():
    file = request.files.get('file')
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)
    try:
        out_filename = f"converted_pdf_{uuid.uuid4().hex[:8]}.pptx"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        convert_pdf_to_pptx(input_path, out_path)
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/excel_to_pdf', methods=['POST'])
def excel_to_pdf():
    file = request.files.get('file')
    if not file or file.filename == '':
        return jsonify({'error': 'Excel file is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('xlsx'))
    file.save(input_path)
    try:
        out_filename = f"converted_xlsx_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        convert_excel_to_pdf(input_path, out_path)
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/pdf_to_excel', methods=['POST'])
def pdf_to_excel():
    file = request.files.get('file')
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)
    try:
        out_filename = f"converted_pdf_{uuid.uuid4().hex[:8]}.xlsx"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        convert_pdf_to_excel(input_path, out_path)
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/pdf_to_pdfa', methods=['POST'])
def pdf_to_pdfa():
    file = request.files.get('file')
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)
    try:
        out_filename = f"archived_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        try:
            doc = fitz.open(input_path)
            doc.save(out_path, pdfa=True)
            doc.close()
        except Exception:
            shutil.copy(input_path, out_path)
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/repair_pdf', methods=['POST'])
def repair_pdf():
    file = request.files.get('file')
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)
    try:
        out_filename = f"repaired_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        doc = fitz.open(input_path)
        doc.save(out_path, garbage=4, deflate=True, clean=True)
        doc.close()
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/ocr_pdf', methods=['POST'])
def ocr_pdf():
    file = request.files.get('file')
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)
    try:
        out_filename = f"ocr_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        run_ocr_on_pdf(input_path, out_path)
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/pdf_summarize', methods=['POST'])
def pdf_summarize():
    file = request.files.get('file')
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)
    try:
        out_filename = f"summary_{uuid.uuid4().hex[:8]}.txt"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        summarize_pdf_document(input_path, out_path)
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/translate_pdf', methods=['POST'])
def translate_pdf():
    file = request.files.get('file')
    target_lang = request.form.get('target_lang', 'es')
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)
    try:
        out_filename = f"translated_{target_lang}_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        translate_pdf_document(input_path, out_path, target_lang)
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/sign_pdf', methods=['POST'])
def sign_pdf():
    file = request.files.get('file')
    sig_b64 = request.form.get('signature_base64')
    try:
        page_idx = int(request.form.get('page_index', 0))
        left = float(request.form.get('left', 0.0))
        top = float(request.form.get('top', 0.0))
        width = float(request.form.get('width', 150.0))
        height = float(request.form.get('height', 60.0))
    except ValueError:
        return jsonify({'error': 'Invalid placement coordinates.'}), 400
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    if not sig_b64:
        return jsonify({'error': 'Signature image is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)
    try:
        import base64
        img_bytes = base64.b64decode(sig_b64)
        out_filename = f"signed_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        doc = fitz.open(input_path)
        page = doc[page_idx]
        rect = fitz.Rect(left, top, left + width, top + height)
        page.insert_image(rect, stream=img_bytes)
        doc.save(out_path)
        doc.close()
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/<path:slug>')
def dynamic_seo_page(slug):
    slug = slug.strip('/')
    
    # Allow serving static files from root (like style.css, ad-manager.js, images)
    root_file = os.path.join(BASE_DIR, slug)
    print(f"DEBUG: Requesting slug={slug}, root_file={root_file}, exists={os.path.exists(root_file)}")
    
    if os.path.exists(root_file) and os.path.isfile(root_file):
        print(f"DEBUG: Sending from directory {BASE_DIR}")
        return send_from_directory(BASE_DIR, slug)
        
    if not slug.endswith('.html'):
        slug += '.html'
        
    file_path = os.path.join(BASE_DIR, 'seo_pages', slug)
    print(f"DEBUG: Fallback to seo_pages path={file_path}, exists={os.path.exists(file_path)}")
    if os.path.exists(file_path):
        return send_from_directory('seo_pages', slug)
    
    print("DEBUG: Aborting with 404")
    abort(404)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)

